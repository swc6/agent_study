import os
from docling_parse.pdf_parser import DoclingPdfParser
from docx import Document

# ============ 配置区 ============
# 定义要解析的文件名（只需修改这里）
FILE_NAME = "2.docx"  # 支持 .pdf 和 .docx 格式
# =================================


class PdfProcessor:
    """PDF文档处理器"""

    def __init__(self):
        self.parser = DoclingPdfParser()

    def load(self, file_path: str):
        """加载PDF文档"""
        return self.parser.load(file_path)

    def extract_text(self, doc, max_pages: int = 2) -> str:
        """提取文本内容"""
        all_text = []
        total_pages = min(doc.number_of_pages(), max_pages)

        for page_no in range(1, total_pages + 1):
            page = doc.get_page(page_no)
            page_text = self._extract_page_text(page)
            all_text.append(f"--- 第{page_no}页 ---\n{page_text}")

        return "\n\n".join(all_text)

    def _extract_page_text(self, page) -> str:
        """从页面提取文本"""
        text = ""

        if hasattr(page, 'textline_cells') and page.textline_cells:
            for line in page.textline_cells:
                text += line.text + "\n"
        elif hasattr(page, 'word_cells') and page.word_cells:
            for word in page.word_cells:
                text += word.text + " "
        elif hasattr(page, 'char_cells') and page.char_cells:
            for char in page.char_cells:
                text += char.text

        return text

    def extract_structure(self, doc) -> dict:
        """提取文档结构"""
        structure = {
            "pages": doc.number_of_pages(),
            "table_of_contents": []
        }

        try:
            annotations = doc.get_annotations()
            if annotations and annotations.table_of_contents:
                for entry in annotations.table_of_contents[:10]:
                    level = entry.level if entry.level else 0
                    structure["table_of_contents"].append({
                        "level": level,
                        "title": entry.title,
                        "page": entry.page
                    })
        except Exception:
            pass

        return structure


class DocxProcessor:
    """DOCX文档处理器"""

    def load(self, file_path: str):
        """加载DOCX文档"""
        return Document(file_path)

    def extract_text(self, doc, max_items: int = 50) -> str:
        """提取文本内容"""
        all_text = []
        paragraphs = list(doc.paragraphs)[:max_items]

        for i, para in enumerate(paragraphs):
            if para.text.strip():
                all_text.append(para.text)

        return "\n\n".join(all_text)

    def extract_tables(self, doc) -> list:
        """提取表格"""
        tables = []
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append({
                "index": i,
                "rows": len(table.rows),
                "data": table_data
            })
        return tables

    def extract_structure(self, doc) -> dict:
        """提取文档结构"""
        structure = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "headings": []
        }

        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                structure["headings"].append({
                    "level": para.style.name,
                    "text": para.text[:100]
                })

        return structure


def get_processor(file_name: str):
    """根据文件后缀获取对应的处理器"""
    ext = os.path.splitext(file_name)[1].lower()

    processors = {
        '.pdf': PdfProcessor(),
        '.docx': DocxProcessor()
    }

    processor = processors.get(ext)
    if not processor:
        raise ValueError(f"不支持的文件格式: {ext}，支持的格式: .pdf, .docx")

    return processor


def main():
    print("=== 文档解析练习 ===\n")
    print(f"配置文件名: {FILE_NAME}\n")

    if not os.path.exists(FILE_NAME):
        print(f"错误: 文件 '{FILE_NAME}' 不存在")
        print("提示: 请确保文件在当前目录下")
        return

    try:
        # 获取对应的处理器
        processor = get_processor(FILE_NAME)
        print(f"使用处理器: {processor.__class__.__name__}\n")

        # 加载文档
        print("1. 加载文档...")
        doc = processor.load(FILE_NAME)
        print("   ✓ 文档加载成功")

        # 提取文本
        print("\n2. 提取文本内容...")
        text = processor.extract_text(doc, max_items=50)
        print(f"   提取文本长度: {len(text)} 字符")
        print(f"\n   文本预览:\n{text[:800]}...")

        # 提取结构
        print("\n3. 提取文档结构...")
        structure = processor.extract_structure(doc)
        print(f"   文档结构: {structure}")

        # 如果是DOCX，额外提取表格
        if isinstance(processor, DocxProcessor):
            print("\n4. 提取表格...")
            tables = processor.extract_tables(doc)
            print(f"   发现 {len(tables)} 个表格")
            for table in tables[:2]:
                print(f"\n   表格 {table['index'] + 1} ({table['rows']}行):")
                for row in table['data'][:3]:
                    print(f"     {row}")

        print("\n✓ 文档解析完成!")

    except Exception as e:
        print(f"\n错误: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()